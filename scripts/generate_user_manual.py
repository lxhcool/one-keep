from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SNAP_DIR = ROOT / "snap"
OUT_DIR = ROOT / "软著申请材料-厘清V1.0"
OUT_PATH = OUT_DIR / "厘清记账软件V1.0_用户操作手册_按模板.docx"
PDF_PATH = OUT_DIR / "厘清记账软件V1.0_用户操作手册_按模板.pdf"

SOFTWARE = "厘清记账软件 V1.0"
PDF_FONT = "ArialUnicode"
PDF_FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"


PAGES = [
    {
        "title": "账号登录",
        "image": "01_账号登录页.jpg",
        "paragraphs": [
            "用户打开厘清记账软件后进入账号登录界面。",
            "用户在当前界面输入账号和密码，勾选同意用户协议和隐私政策后，点击【登录】按钮完成登录操作。",
            "如尚未注册账号，可点击界面上的【注册】入口进入账号注册页面。",
        ],
    },
    {
        "title": "账号注册",
        "image": "02_账号注册页.jpg",
        "paragraphs": [
            "用户点击登录页中的【注册】入口后进入账号注册界面。",
            "用户按照页面提示输入用户名、邮箱、密码、昵称等信息，勾选同意相关协议后，点击【注册并登录】按钮完成账号创建。",
            "注册完成后，系统进入软件主界面。",
        ],
    },
    {
        "title": "隐私政策",
        "image": "03_隐私政策.jpg",
        "paragraphs": [
            "用户可在登录或关于页面查看隐私政策内容。",
            "当前页面展示软件关于信息收集、使用、存储和保护的说明，用户可上下滑动查看完整内容。",
            "用户阅读后可点击左上角返回按钮退出当前页面。",
        ],
    },
    {
        "title": "首页",
        "image": "04_主页.jpg",
        "paragraphs": [
            "用户登录成功后进入首页。",
            "首页展示当前用户昵称、本月收入、本月支出、本月结余及最近记录等信息。",
            "用户可点击底部中间的【+】按钮快速新增一笔记账记录，也可通过底部导航切换至统计、账单和我的页面。",
        ],
    },
    {
        "title": "新增记账",
        "image": "05_新增记账.jpg",
        "paragraphs": [
            "用户点击首页或底部导航栏中的【+】按钮后进入新增记账界面。",
            "用户可选择支出或收入类型，输入账单金额，并选择对应分类。",
            "用户填写完成后点击【确认】按钮，系统保存该笔账单并同步更新首页、账单和统计数据。",
        ],
    },
    {
        "title": "编辑账单",
        "image": "06_编辑账单.jpg",
        "paragraphs": [
            "用户在首页最近记录或账单列表中点击某一条账单，可查看账单详情。",
            "当前界面展示账单金额、分类、交易日期、商家名称和交易备注等信息。",
            "用户点击【编辑】按钮可修改账单内容，点击【删除】按钮可删除该条账单。",
        ],
    },
    {
        "title": "删除确认",
        "image": "07_删除确认.jpg",
        "paragraphs": [
            "用户在账单详情中点击【删除】按钮后，系统弹出删除确认窗口。",
            "用户点击【取消】按钮可放弃删除操作；点击【删除】按钮后，该账单记录将从系统中移除。",
            "删除后的账单不再参与首页汇总、账单列表和统计分析。",
        ],
    },
    {
        "title": "账单列表",
        "image": "08_账单页.jpg",
        "paragraphs": [
            "用户点击底部导航中的【账单】按钮进入账单页面。",
            "账单页面按月份和日期展示交易记录，并显示对应交易名称、时间、分类和金额。",
            "用户可点击列表中的账单记录查看详情或进行编辑。",
        ],
    },
    {
        "title": "账单筛选",
        "image": "09_账单筛选.jpg",
        "paragraphs": [
            "用户可在账单页面切换查看全部、支出或收入记录。",
            "点击月份选择区域后，系统弹出月份筛选面板。",
            "用户选择对应年份和月份后，账单列表会展示所选月份的交易数据。",
        ],
    },
    {
        "title": "统计分析",
        "image": "10_统计页.jpg",
        "paragraphs": [
            "用户点击底部导航中的【统计】按钮进入统计页面。",
            "统计页面展示指定月份的收支趋势、收入支出切换、时间维度筛选和分类排行。",
            "用户可通过统计图和排行信息了解个人收支变化及主要消费分类。",
        ],
    },
    {
        "title": "我的页面",
        "image": "11_我的页面.jpg",
        "paragraphs": [
            "用户点击底部导航中的【我的】按钮进入个人中心页面。",
            "当前页面展示用户头像、昵称、本月收入、本月支出以及功能设置入口。",
            "用户可在此进入头部背景、深色模式、分类管理、账号管理和关于系统等功能。",
        ],
    },
    {
        "title": "分类管理",
        "image": "12_分类管理.jpg",
        "paragraphs": [
            "用户在我的页面点击【分类管理】后进入分类管理界面。",
            "分类管理页面分为支出类别和收入类别，展示系统内置和用户自定义的分类。",
            "用户可点击添加按钮新增分类，也可进入已有分类进行调整。",
        ],
    },
    {
        "title": "添加分类",
        "image": "13_添加编辑分类.jpg",
        "paragraphs": [
            "用户点击分类管理中的添加按钮后进入添加分类界面。",
            "用户可选择分类图标并输入分类名称。",
            "填写完成后点击【确认】按钮，系统保存该分类并显示在分类列表中。",
        ],
    },
    {
        "title": "更换头部背景",
        "image": "14_更换头部背景.jpg",
        "paragraphs": [
            "用户在我的页面点击【头部背景】后，可进入头部背景更换窗口。",
            "系统提供预设背景供用户选择，也支持从相册选择图片。",
            "用户选择后，个人中心顶部背景将更新为新的显示效果。",
        ],
    },
    {
        "title": "更改头像",
        "image": "15_更改头像.jpg",
        "paragraphs": [
            "用户点击个人中心头像区域后，可进入头像更改窗口。",
            "用户可从相册选择图片作为头像，也可将头像重置为默认头像。",
            "头像设置完成后，会在个人中心页面展示。",
        ],
    },
    {
        "title": "关于系统",
        "image": "16_关于系统.jpg",
        "paragraphs": [
            "用户在我的页面点击【关于系统】后进入关于系统界面。",
            "当前页面展示软件名称、版本号、用户协议、隐私政策及相关说明。",
            "用户可点击对应入口查看协议或隐私政策内容。",
        ],
    },
    {
        "title": "退出登录",
        "image": "17_退出登录.jpg",
        "paragraphs": [
            "用户在我的页面进入账号管理后，可进行退出登录操作。",
            "点击【退出登录】按钮后，系统退出当前账号并返回登录界面。",
            "退出后再次使用软件，需要重新输入账号和密码进行登录。",
        ],
    },
    {
        "title": "注销账号",
        "image": "18_注销账号.jpg",
        "paragraphs": [
            "用户在账号管理中点击【注销账号】后，系统弹出注销确认窗口。",
            "窗口会提示注销账号将清除账号资料、账单记录、分类和个人设置等信息。",
            "用户确认后执行注销操作；如不继续注销，可点击【取消】返回。",
        ],
    },
]


def set_east_asian_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_border_bottom(paragraph, color="000000", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.6))
    run = p.add_run(SOFTWARE)
    set_east_asian_font(run, "宋体")
    run.font.size = Pt(10.5)
    p.add_run("\t")
    add_page_field(p)
    set_paragraph_border_bottom(p)


def add_cover(doc):
    for _ in range(7):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE)
    set_east_asian_font(run, "宋体")
    run.font.size = Pt(24)
    run.bold = True
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用户操作手册")
    set_east_asian_font(run, "宋体")
    run.font.size = Pt(22)
    run.bold = True
    doc.add_page_break()


def add_manual_page(doc, item):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run(item["title"])
    set_east_asian_font(run, "宋体")
    run.font.size = Pt(15)
    run.bold = True

    for text in item["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        set_east_asian_font(run, "宋体")
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    img_path = SNAP_DIR / item["image"]
    p.add_run().add_picture(str(img_path), height=Cm(15.2))


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    for idx, page in enumerate(PAGES):
        add_manual_page(doc, page)
        if idx != len(PAGES) - 1:
            doc.add_page_break()

    doc.save(OUT_PATH)
    build_pdf()
    print(OUT_PATH)
    print(PDF_PATH)


def draw_wrapped_text(c, text, x, y, max_width, font_name, font_size, line_gap=7):
    chars = list(text)
    lines = []
    line = ""
    for char in chars:
        trial = line + char
        if c.stringWidth(trial, font_name, font_size) <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = char
    if line:
        lines.append(line)

    for line in lines:
        c.drawString(x, y, line)
        y -= font_size + line_gap
    return y


def draw_header(c, page_no, width, height):
    left = 3.0 * cm
    right = width - 2.6 * cm
    y = height - 1.18 * cm
    c.setFont(PDF_FONT, 10.5)
    c.drawString(left, y, SOFTWARE)
    c.line(left, y - 0.14 * cm, right - 0.5 * cm, y - 0.14 * cm)
    c.drawRightString(right, y, str(page_no))


def draw_cover(c, page_no, width, height):
    draw_header(c, page_no, width, height)
    c.setFont(PDF_FONT, 24)
    c.drawCentredString(width / 2, height - 8.7 * cm, SOFTWARE)
    c.setFont(PDF_FONT, 22)
    c.drawCentredString(width / 2, height - 11.7 * cm, "用户操作手册")


def draw_content_page(c, page_no, item, width, height):
    left = 3.0 * cm
    right = width - 2.6 * cm
    draw_header(c, page_no, width, height)

    y = height - 3.05 * cm
    c.setFont(PDF_FONT, 15)
    c.drawString(left, y, item["title"])
    y -= 1.15 * cm

    c.setFont(PDF_FONT, 12)
    for text in item["paragraphs"]:
        y = draw_wrapped_text(
            c,
            "    " + text,
            left,
            y,
            right - left,
            PDF_FONT,
            12,
            line_gap=6,
        )
        y -= 0.16 * cm

    img_path = SNAP_DIR / item["image"]
    with Image.open(img_path) as img:
        iw, ih = img.size
    max_h = 15.2 * cm
    max_w = 7.8 * cm
    scale = min(max_w / iw, max_h / ih)
    draw_w = iw * scale
    draw_h = ih * scale
    img_x = (width - draw_w) / 2
    img_y = max(1.45 * cm, y - 0.35 * cm - draw_h)
    c.drawImage(str(img_path), img_x, img_y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask="auto")


def build_pdf():
    pdfmetrics.registerFont(TTFont(PDF_FONT, PDF_FONT_PATH))
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    width, height = A4
    draw_cover(c, 1, width, height)
    c.showPage()

    for idx, item in enumerate(PAGES, start=2):
        draw_content_page(c, idx, item, width, height)
        if idx != len(PAGES) + 1:
            c.showPage()

    c.save()


if __name__ == "__main__":
    main()
